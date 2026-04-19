import streamlit as st
from groq import Groq
from docx import Document
from io import BytesIO

# --- API AYARI ---
# Streamlit 'Sırlar' (Secrets) kısmına GROQ_API_KEY olarak eklediğin anahtarı çeker
client = Groq(api_key=st.secrets["GROQ_API_KEY"])

# --- SAYFA AYARLARI ---
st.set_page_config(page_title="Enes Altun | Yapay Zeka CV Stüdyosu", page_icon="🚀", layout="wide")

# --- CSS EFEKTLERİ ---
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .stButton>button { 
        width: 100%; 
        border-radius: 25px; 
        height: 3.5em; 
        background-color: #007bff; 
        color: white; 
        font-weight: bold;
        transition: 0.3s;
    }
    .stButton>button:hover { background-color: #0056b3; border: 2px solid white; }
    .stTextInput>div>div>input { border-radius: 12px; border: 1px solid #ddd; }
    .stTextArea>div>div>textarea { border-radius: 12px; border: 1px solid #ddd; }
    </style>
    """, unsafe_allow_html=True)

# --- UYGULAMA BAŞLIĞI ---
st.title("📄 Yapay Zeka Destekli CV Stüdyosu")
st.info("Kanka bilgilerini gir, profesyonel CV'ni saniyeler içinde Word olarak indirelim!")

# --- FORM KUTUCUKLARI ---
with st.container():
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("👤 Kişisel Bilgiler")
        ad_soyad = st.text_input("Adınız Soyadınız", placeholder="Örn: Enes Altun")
        eposta = st.text_input("E-posta Adresiniz", placeholder="enes@ornek.com")
        telefon = st.text_input("Telefon Numaranız", placeholder="0555 xxx xx xx")

    with col2:
        st.subheader("🎓 Eğitim ve Deneyim")
        egitim = st.text_area("Eğitim Bilgileriniz", placeholder="Okul adı, bölüm ve mezuniyet yılı...")
        deneyim = st.text_area("İş/Proje Deneyimleriniz", placeholder="Daha önce neler yaptın? Stajlar, projeler...")
        yetenekler = st.text_input("Yetenekleriniz", placeholder="Python, Office, Video Kurgu vb.")

# --- CV OLUŞTURMA İŞLEMİ ---
if st.button("Profesyonel CV Oluştur ✨"):
    if ad_soyad and eposta:
        try:
            with st.spinner('🚀 Yapay zeka CV\'ni hazırlıyor, lütfen bekle...'):
                # Yapay zekaya giden talimat
                prompt = f"""
                Lütfen aşağıdaki bilgilere sahip bir kişi için profesyonel, etkileyici ve modern bir CV içeriği oluştur.
                Dili profesyonel olsun.
                
                Ad Soyad: {ad_soyad}
                E-posta: {eposta}
                Telefon: {telefon}
                Eğitim: {egitim}
                Deneyim: {deneyim}
                Yetenekler: {yetenekler}
                """
                
                # MODEL GÜNCELLENDİ: llama-3.3-70b-versatile
                chat_completion = client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    model="llama-3.3-70b-versatile",
                )
                
                cv_metni = chat_completion.choices[0].message.content
                
                st.success("✅ CV Başarıyla Hazırlandı!")
                st.markdown("---")
                st.subheader("📋 CV Taslağın")
                st.write(cv_metni)
                
                # Word Dosyası Oluşturma
                doc = Document()
                doc.add_heading(ad_soyad, 0)
                doc.add_paragraph(cv_metni)
                
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(
                    label="📥 Word Dosyası Olarak İndir",
                    data=buffer,
                    file_name=f"{ad_soyad}_CV.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"Eyvah! Bir hata oluştu kanka: {e}")
            st.info("İpucu: 'Sırlar' kısmındaki API anahtarını kontrol etmeyi unutma!")
    else:
        st.warning("Kanka en azından adını ve e-postanı yazman lazım ki sistem çalışsın!")
